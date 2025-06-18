import pandas as pd
from sqlalchemy.orm import Session
import os # For ensuring export directory exists

# CRUD functions
from app.crud import employee_crud
from app.crud import hardware_crud
from app.crud import license_crud
from app.crud import web_access_crud

# Pydantic Schemas for data transformation/selection before export
# (already imported for Excel export, ensure they are available)
from app.models import employee_schemas, hardware_schemas, license_schemas, web_access_schemas

# Imports for Word document generation
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH # For text alignment if needed

# Ensure the export directory exists
EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

def generate_excel_report(db: Session) -> str:
    output_path = os.path.join(EXPORT_DIR, "reporte_activos_tic.xlsx")

    # 1. Fetch data using CRUD functions
    employees_db = employee_crud.get_employees(db, limit=-1) # Get all
    hardware_db = hardware_crud.get_hardware_items(db, limit=-1)
    licenses_db = license_crud.get_licenses(db, limit=-1)
    web_accesses_db = web_access_crud.get_all_web_accesses(db, limit=-1)

    # 2. Convert ORM objects to Pydantic schemas for consistent field selection
    # This also helps in preparing data for DataFrame conversion, especially if using .dict()

    # Employees: Select basic fields, potentially expand with related data if safe & simple
    employees_data = [employee_schemas.Employee.from_orm(emp).dict(exclude_none=True) for emp in employees_db]

    # Hardware: Select fields, potentially map employee_id to name if required (more complex)
    # For now, keeping it simple with IDs.
    hardware_data = [hardware_schemas.Hardware.from_orm(hw).dict(exclude_none=True) for hw in hardware_db]

    # Licenses: Similar to hardware
    licenses_data = [license_schemas.License.from_orm(lic).dict(exclude_none=True) for lic in licenses_db]

    # Web Accesses: CRITICAL: Exclude sensitive information like hashed passwords
    # The WebAccess schema is already designed to exclude passwords from responses.
    # We'll manually select fields here for the DataFrame to be explicit.
    web_accesses_export_data = []
    for wa in web_accesses_db:
        web_accesses_export_data.append({
            "id": wa.id,
            "nombre_servicio": wa.nombre_servicio,
            "url": wa.url,
            "usuario": wa.usuario,
            "descripcion": wa.descripcion,
            "notas_adicionales": wa.notas_adicionales,
            "employee_id": wa.employee_id
            # Add employee name here with a lookup if desired
        })

    # 3. Convert lists of dicts/Pydantic objects to Pandas DataFrames
    df_employees = pd.DataFrame(employees_data)
    df_hardware = pd.DataFrame(hardware_data)
    df_licenses = pd.DataFrame(licenses_data)
    df_web_accesses = pd.DataFrame(web_accesses_export_data)

    # Clean up DataFrames: remove ORM relationship objects if they were included by .from_orm()
    # and not handled by .dict(exclude={...}) or schema definition.
    # For example, if Employee schema included full Hardware objects:
    if 'hardware_items' in df_employees.columns:
      df_employees = df_employees.drop(columns=['hardware_items'])
    if 'licenses' in df_employees.columns: # if full license objects were included
      df_employees = df_employees.drop(columns=['licenses'])
    if 'web_accesses' in df_employees.columns:
      df_employees = df_employees.drop(columns=['web_accesses'])

    if 'employee' in df_hardware.columns: # if full employee object was included
        df_hardware = df_hardware.drop(columns=['employee'])
    if 'licenses' in df_hardware.columns:
        df_hardware = df_hardware.drop(columns=['licenses'])

    if 'employee' in df_licenses.columns:
        df_licenses = df_licenses.drop(columns=['employee'])
    if 'hardware' in df_licenses.columns:
        df_licenses = df_licenses.drop(columns=['hardware'])

    # df_web_accesses already manually constructed to be flat

    # 4. Write DataFrames to Excel sheets
    try:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df_employees.to_excel(writer, sheet_name='Empleados', index=False)
            df_hardware.to_excel(writer, sheet_name='Hardware', index=False)
            df_licenses.to_excel(writer, sheet_name='Licencias', index=False)
            df_web_accesses.to_excel(writer, sheet_name='AccesosWeb', index=False)
        # writer.save() is called automatically when exiting the 'with' block if using pandas >= 1.2.0
        # For older pandas, writer.save() might be needed explicitly if not using 'with'.
    except Exception as e:
        # Handle potential errors during Excel writing, e.g., file permissions
        print(f"Error writing Excel file: {e}")
        # Consider raising an exception or returning an error status
        raise

    return output_path


def generate_word_acta(db: Session, employee_id: int) -> Optional[str]:
    employee = employee_crud.get_employee(db, employee_id=employee_id)
    if not employee:
        return None # Or raise HTTPException if preferred for API level

    document = Document()

    # Set margins (optional) - example: 2cm for all
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title = document.add_heading('Acta de Entrega de Activos TIC', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center align the title

    # Employee Information
    document.add_heading('Datos del Empleado', level=2)
    document.add_paragraph(f'Nombre: {employee.nombre}')
    document.add_paragraph(f'Identificación: {employee.identificacion}')
    document.add_paragraph(f'Cargo: {employee.cargo}')
    if employee.departamento:
        document.add_paragraph(f'Departamento: {employee.departamento}')
    document.add_paragraph() # Add a blank line for spacing

    # Hardware Assigned
    document.add_heading('Hardware Asignado', level=2)
    if employee.hardware_items:
        table = document.add_table(rows=1, cols=4) # Added 'Modelo'
        table.style = 'Table Grid' # Apply a grid style
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tipo Equipo'
        hdr_cells[1].text = 'Marca'
        hdr_cells[2].text = 'Modelo'
        hdr_cells[3].text = 'Serial'
        for item in employee.hardware_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.tipo_equipo
            row_cells[1].text = item.marca
            row_cells[2].text = item.modelo if item.modelo else ""
            row_cells[3].text = item.serial
    else:
        document.add_paragraph("No hay hardware asignado.")
    document.add_paragraph()

    # Licenses Assigned
    document.add_heading('Licencias Asignadas', level=2)
    if employee.licenses:
        table = document.add_table(rows=1, cols=3) # Software, Tipo, Clave (optional)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Software'
        hdr_cells[1].text = 'Tipo Licencia'
        hdr_cells[2].text = 'Clave Producto (Parcial)' # Advise on not printing full keys
        for lic in employee.licenses:
            row_cells = table.add_row().cells
            row_cells[0].text = lic.software_nombre
            row_cells[1].text = lic.tipo_licencia
            # Displaying partial or no product key for security is advisable
            clave_display = f"********{lic.clave_producto[-4:]}" if lic.clave_producto and len(lic.clave_producto) > 4 else "N/A"
            row_cells[2].text = clave_display if lic.clave_producto else "N/A"
    else:
        document.add_paragraph("No hay licencias asignadas.")
    document.add_paragraph()

    # Web Accesses - Generally not included in a physical delivery act.
    # If needed, add a section similar to above, carefully considering password info.
    # For this act, we assume Web Accesses are managed separately or not part of this physical document.

    # Acceptance and Signatures
    document.add_heading('Confirmación y Firmas', level=2)
    document.add_paragraph(
        "Declaro haber recibido los activos TIC listados en esta acta en buen estado y funcionamiento, "
        "y me comprometo a utilizarlos de acuerdo con las políticas de la empresa."
    )
    document.add_paragraph()

    # Using a table for cleaner signature layout
    sig_table = document.add_table(rows=1, cols=2)
    # Remove table borders if desired by setting style or iterating cells

    # Cell for "Entrega"
    cell_entrega = sig_table.cell(0, 0)
    p_entrega = cell_entrega.add_paragraph()
    p_entrega.add_run("\n\n\n_________________________\n").bold = False
    p_entrega.add_run("Firma Responsable TI (Entrega)").bold = True
    # Add more details if needed: Nombre, Cargo, Fecha

    # Cell for "Recibe"
    cell_recibe = sig_table.cell(0, 1)
    p_recibe = cell_recibe.add_paragraph()
    p_recibe.add_run("\n\n\n_________________________\n").bold = False
    p_recibe.add_run(f"Firma Empleado ({employee.nombre})").bold = True
    # Add more details: Identificación, Cargo, Fecha

    # Define output path
    safe_identificacion = "".join(c if c.isalnum() else "_" for c in employee.identificacion)
    file_path = os.path.join(EXPORT_DIR, f"acta_entrega_{safe_identificacion}.docx")

    try:
        document.save(file_path)
    except Exception as e:
        print(f"Error saving Word document: {e}")
        # Consider raising an exception or returning an error status
        raise

    return file_path
